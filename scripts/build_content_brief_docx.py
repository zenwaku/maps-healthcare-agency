from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "MAPS-Landing-Page-Content-Brief.md"
OUTPUT = ROOT / "docs" / "MAPS-Landing-Page-Content-Brief-Revised.docx"

NAVY = RGBColor(15, 23, 42)
TEAL = RGBColor(15, 118, 110)
BLUE = RGBColor(37, 99, 235)
MUTED = RGBColor(71, 85, 105)
ORANGE = RGBColor(249, 115, 22)
LIGHT = "EEF9F6"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, TEAL, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = NAVY
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_rich_text(paragraph, text, color=NAVY):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = paragraph.add_run(clean)
        set_font(run, size=11, color=color, bold=is_bold)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("MAPS")
    set_font(run, size=12, color=ORANGE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Landing Page Content Brief")
    set_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("Medical Advance Portfolio Scientific")
    set_font(run, size=16, color=TEAL, bold=True)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(30)
    run = tagline.add_run("Doctor-Led Digital Marketing for Healthcare Growth")
    set_font(run, size=12, color=MUTED, italic=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.left_indent = Inches(0.8)
    note.paragraph_format.right_indent = Inches(0.8)
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(0)
    set_cell_shading(note, LIGHT)
    run = note.add_run("Master copy reference for website editing, review, and future campaign alignment.")
    set_font(run, size=10.5, color=NAVY, bold=True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(42)
    run = date.add_run("19 June 2026")
    set_font(run, size=10, color=MUTED)

    doc.add_page_break()


def add_header_footer(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("MAPS  |  Landing Page Content Brief")
    set_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("Medical Advance Portfolio Scientific  |  ")
    set_font(run, size=9, color=MUTED)
    add_page_field(paragraph)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    skipped_title = False
    for line in lines:
        text = line.strip()
        if not text:
            continue
        if text.startswith("# ") and not skipped_title:
            skipped_title = True
            continue
        if text.startswith("## "):
            doc.add_paragraph(text[3:], style="Heading 1")
            continue
        if text.startswith("### "):
            doc.add_paragraph(text[4:], style="Heading 2")
            continue
        if text.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, text[2:])
            continue
        if re.match(r"^\d+\.\s", text):
            paragraph = doc.add_paragraph(style="List Number")
            add_rich_text(paragraph, re.sub(r"^\d+\.\s+", "", text))
            continue

        paragraph = doc.add_paragraph()
        add_rich_text(paragraph, text)

    props = doc.core_properties
    props.title = "MAPS Landing Page Content Brief"
    props.subject = "Master website copy reference"
    props.author = "MAPS - Medical Advance Portfolio Scientific"
    props.keywords = "MAPS, healthcare digital marketing, landing page, content brief"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
